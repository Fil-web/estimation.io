from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
from xml.sax.saxutils import escape

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont

FONT_PATH = Path("/Library/Fonts/Arial Unicode.ttf")
PAGE_SIZE = (1240, 1754)
MARGIN = 70


def _load_font(size, bold=False):
    if FONT_PATH.exists():
        return ImageFont.truetype(str(FONT_PATH), size=size)
    return ImageFont.load_default()


def _wrap_text(draw, text, font, max_width):
    words = str(text or "").split()
    if not words:
        return [""]

    lines = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def _new_pdf_page():
    image = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(image)
    return image, draw


def _save_pdf(images):
    buffer = BytesIO()
    first, *rest = images
    first.save(buffer, format="PDF", resolution=150.0, save_all=True, append_images=rest)
    return buffer.getvalue()


def _column_letter(index):
    result = ""
    while index > 0:
        index, remainder = divmod(index - 1, 26)
        result = chr(65 + remainder) + result
    return result


def _sheet_xml(rows):
    xml_rows = []
    for row_index, row in enumerate(rows, start=1):
        cells = []
        for col_index, value in enumerate(row, start=1):
            if value in (None, ""):
                continue
            cell_ref = f"{_column_letter(col_index)}{row_index}"
            text = escape(str(value))
            cells.append(
                f'<c r="{cell_ref}" t="inlineStr"><is><t xml:space="preserve">{text}</t></is></c>'
            )
        xml_rows.append(f'<row r="{row_index}">{"".join(cells)}</row>')

    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f'<sheetData>{"".join(xml_rows)}</sheetData>'
        "</worksheet>"
    )


def _build_xlsx(sheet_name, rows):
    workbook_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<sheets><sheet name="{escape(sheet_name)}" sheetId="1" r:id="rId1"/></sheets>'
        "</workbook>"
    )
    workbook_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
        'Target="worksheets/sheet1.xml"/>'
        "</Relationships>"
    )
    root_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="xl/workbook.xml"/>'
        "</Relationships>"
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/xl/workbook.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
        '<Override PartName="/xl/worksheets/sheet1.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        "</Types>"
    )

    buffer = BytesIO()
    with ZipFile(buffer, "w", compression=ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", root_rels)
        archive.writestr("xl/workbook.xml", workbook_xml)
        archive.writestr("xl/_rels/workbook.xml.rels", workbook_rels)
        archive.writestr("xl/worksheets/sheet1.xml", _sheet_xml(rows))
    return buffer.getvalue()


def _draw_table(draw, start_y, headers, rows, col_widths, title_lines):
    font = _load_font(21)
    small_font = _load_font(18)
    title_font = _load_font(28)
    bold_font = _load_font(21)

    x0 = MARGIN
    y = start_y

    for line in title_lines:
        bbox = draw.textbbox((0, 0), line, font=title_font)
        width = bbox[2] - bbox[0]
        draw.text(((PAGE_SIZE[0] - width) / 2, y), line, fill="black", font=title_font)
        y += 42

    y += 18

    row_height = 44
    x = x0
    for index, header in enumerate(headers):
        width = col_widths[index]
        draw.rectangle([x, y, x + width, y + row_height], outline="black", width=2)
        lines = _wrap_text(draw, header, bold_font, width - 16)
        text_y = y + 8
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=bold_font)
            text_x = x + (width - (bbox[2] - bbox[0])) / 2
            draw.text((text_x, text_y), line, fill="black", font=bold_font)
            text_y += 20
        x += width
    y += row_height

    for row in rows:
        prepared = []
        max_lines = 1
        for index, value in enumerate(row):
            lines = _wrap_text(draw, value, small_font, col_widths[index] - 14)
            prepared.append(lines)
            max_lines = max(max_lines, len(lines))

        current_height = max(40, 12 + max_lines * 22)
        x = x0
        for index, lines in enumerate(prepared):
            width = col_widths[index]
            draw.rectangle([x, y, x + width, y + current_height], outline="black", width=1)
            text_y = y + 8
            for line in lines:
                draw.text((x + 6, text_y), line, fill="black", font=small_font)
                text_y += 22
            x += width
        y += current_height

    return y


def build_order_docx(period, payments_df):
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10)

    header = document.add_paragraph("Читинский институт Байкальского государственного университета")
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].bold = True
    meta = document.add_paragraph(f"Проект документа от ____________    Период: {period}")
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_heading("Проект приказа", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Обобщенные сведения за период {period}")

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["№", "ФИО", "Должность", "Кафедра", "Итоговые баллы"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header

    for index, row in payments_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(index + 1)
        cells[1].text = str(row["full_name"])
        cells[2].text = str(row["position"] or "-")
        cells[3].text = str(row["department"])
        cells[4].text = f"{row['total_points']:.2f}"

    document.add_paragraph()
    document.add_paragraph("Подготовил: ____________________")
    document.add_paragraph("Согласовано: ____________________")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_order_pdf(period, payments_df):
    image, draw = _new_pdf_page()
    rows = []
    for index, row in payments_df.iterrows():
        rows.append(
            [
                str(index + 1),
                str(row["full_name"]),
                str(row["position"] or "-"),
                str(row["department"]),
                f"{row['total_points']:.2f}",
            ]
        )

    _draw_table(
        draw,
        MARGIN,
        ["№", "ФИО", "Должность", "Кафедра", "Итоговые баллы"],
        rows,
        [60, 330, 220, 320, 180],
        ["Проект приказа", f"Обобщенные сведения за период {period}"],
    )
    return _save_pdf([image])


def build_order_excel(period, payments_df):
    export_df = payments_df[["full_name", "position", "department", "total_points"]].copy()
    export_df.columns = ["ФИО", "Должность", "Кафедра", "Итоговые баллы"]
    export_df.insert(0, "№", range(1, len(export_df) + 1))
    rows = [
        ["Читинский институт Байкальского государственного университета"],
        ["Проект приказа"],
        [f"Обобщенные сведения за период {period}"],
        [""],
        export_df.columns.tolist(),
        *export_df.fillna("").values.tolist(),
        [""],
        ["Подготовил: ____________________"],
        ["Согласовано: ____________________"],
    ]
    return _build_xlsx("Приказ", rows)


def build_service_note_docx(context):
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    document.styles["Normal"].font.name = "Times New Roman"
    document.styles["Normal"].font.size = Pt(9)

    meta = document.add_paragraph(
        f"Читинский институт Байкальского государственного университета\n"
        f"Служебная записка № {context['note_number']} от {context['note_date']}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = document.add_paragraph("Первому заместителю директора\nЧИ ФГБОУ ВО «БГУ»\nН.В. Раевскому")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph()
    h = document.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("СЛУЖЕБНАЯ ЗАПИСКА").bold = True

    subtitle = document.add_paragraph(
        f"Обобщенные сведения о выполнении преподавателями кафедры {context['department_name']} "
        f"критериев эффективности научно-образовательной деятельности для расчета стимулирующих выплат "
        f"за {context['period']} учебный год"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=2, cols=16)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    top = table.rows[0].cells
    second = table.rows[1].cells
    top[0].text = "№"
    top[1].text = "ФИО преподавателя кафедры"
    top[2].text = "Должность"
    labels = [
        "Первый показатель",
        "Второй показатель",
        "Третий показатель",
        "Четвертый показатель",
        "Пятый показатель",
        "Шестой показатель",
    ]
    col = 3
    for label in labels:
        top[col].text = label
        top[col].merge(top[col + 1])
        second[col].text = "Отметка о выполнении, комментарии"
        second[col + 1].text = "Количество баллов"
        col += 2
    top[15].text = "Итоговая сумма баллов"

    top[0].merge(second[0])
    top[1].merge(second[1])
    top[2].merge(second[2])
    top[15].merge(second[15])

    for index, teacher in enumerate(context["teachers"], start=1):
        chunks = [teacher["items"][i:i + 6] for i in range(0, len(teacher["items"]), 6)] or [[]]
        start_row_index = len(table.rows)
        for chunk in chunks:
            row_cells = table.add_row().cells
            row_cells[0].text = str(index)
            row_cells[1].text = teacher["full_name"]
            row_cells[2].text = teacher["position"] or "-"
            c = 3
            for item in chunk:
                comment = item["teacher_comment"] or item["criterion_name"]
                row_cells[c].text = f"{item['code']}\n{comment}\nКоличество: {item['quantity']}"
                row_cells[c + 1].text = f"{float(item['claimed_score']):.2f}"
                c += 2
            while c < 15:
                row_cells[c].text = ""
                row_cells[c + 1].text = ""
                c += 2
            row_cells[15].text = f"{teacher['total_points']:.2f}"

        if len(chunks) > 1:
            end_row_index = len(table.rows) - 1
            for col_index in [0, 1, 2, 15]:
                table.cell(start_row_index, col_index).merge(table.cell(end_row_index, col_index))

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.add_paragraph()
    document.add_paragraph(f"Заведующий кафедрой ____________ /{context['head_name'] or '____________________'}/")
    document.add_paragraph("Согласовано:")
    document.add_paragraph("Начальник отдела учебно-методического и информационного обеспечения ____________")
    document.add_paragraph("Главный специалист по кадрам ____________")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_service_note_pdf(context):
    image, draw = _new_pdf_page()
    rows = []
    for index, teacher in enumerate(context["teachers"], start=1):
        chunks = [teacher["items"][i:i + 6] for i in range(0, len(teacher["items"]), 6)] or [[]]
        first = True
        for chunk in chunks:
            row = [
                str(index) if first else "",
                teacher["full_name"] if first else "",
                teacher["position"] if first else "",
            ]
            for item in chunk:
                comment = item["teacher_comment"] or item["criterion_name"]
                row.append(f"{item['code']}\n{comment}\nКоличество: {item['quantity']}")
                row.append(f"{float(item['claimed_score']):.2f}")
            while len(row) < 15:
                row.extend(["", ""])
            row.append(f"{teacher['total_points']:.2f}" if first else "")
            rows.append(row[:16])
            first = False

    headers = ["№", "ФИО", "Должность"]
    for label in ["1", "Баллы", "2", "Баллы", "3", "Баллы", "4", "Баллы", "5", "Баллы", "6", "Баллы"]:
        headers.append(label)
    headers.append("Итог")

    _draw_table(
        draw,
        30,
        headers,
        rows or [[""] * 16],
        [40, 150, 110, 110, 60, 110, 60, 110, 60, 110, 60, 110, 60, 110, 60, 80],
        [
            f"Служебная записка № {context['note_number']} от {context['note_date']}",
            "СЛУЖЕБНАЯ ЗАПИСКА",
            f"Кафедра: {context['department_name']}",
            f"Период: {context['period']} учебный год",
        ],
    )
    return _save_pdf([image])


def build_service_note_excel(context):
    rows = []
    for index, teacher in enumerate(context["teachers"], start=1):
        chunks = [teacher["items"][i:i + 6] for i in range(0, len(teacher["items"]), 6)] or [[]]
        first = True
        for chunk in chunks:
            row = {
                "№": index if first else "",
                "ФИО преподавателя кафедры": teacher["full_name"] if first else "",
                "Должность": teacher["position"] if first else "",
            }
            for pos in range(6):
                item = chunk[pos] if pos < len(chunk) else None
                row[f"Показатель {pos + 1}: комментарий"] = (
                    f"{item['code']}; {(item['teacher_comment'] or item['criterion_name'])}; Количество: {item['quantity']}"
                    if item
                    else ""
                )
                row[f"Показатель {pos + 1}: баллы"] = f"{float(item['claimed_score']):.2f}" if item else ""
            row["Итоговая сумма баллов"] = f"{teacher['total_points']:.2f}" if first else ""
            rows.append(row)
            first = False

    export_df = pd.DataFrame(rows or [{}])
    excel_rows = [
        ["Читинский институт Байкальского государственного университета"],
        [f"Служебная записка № {context['note_number']} от {context['note_date']}"],
        ["СЛУЖЕБНАЯ ЗАПИСКА"],
        [
            f"Обобщенные сведения по кафедре {context['department_name']} "
            f"за {context['period']} учебный год"
        ],
        [""],
        export_df.columns.tolist(),
        *export_df.fillna("").values.tolist(),
        [""],
        [f"Заведующий кафедрой ____________ /{context['head_name'] or '____________________'}/"],
        ["Согласовано: ____________________"],
    ]
    return _build_xlsx("Записка", excel_rows)
